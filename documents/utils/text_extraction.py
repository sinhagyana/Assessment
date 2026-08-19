import logging
import os

logger = logging.getLogger("documents")


class TextExtractionError(Exception):
    """Raised when text cannot be extracted from a document."""


def extract_text(file_path: str, file_type: str) -> str:
    file_type = file_type.lower()
    try:
        if file_type == ".pdf":
            text = _extract_pdf(file_path)
        elif file_type == ".docx":
            text = _extract_docx(file_path)
        elif file_type == ".txt":
            text = _extract_txt(file_path)
        else:
            raise TextExtractionError(f"Unsupported file type: {file_type}")
    except TextExtractionError:
        raise
    except Exception as exc: 
        logger.exception("Failed to extract text from %s", file_path)
        raise TextExtractionError(f"Error extracting text: {exc}") from exc

    text = (text or "").strip()
    if not text:
        raise TextExtractionError("No extractable text found in document.")
    return text


def _extract_pdf(file_path: str) -> str:
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages.append(page_text)
    return "\n".join(pages)


def _extract_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def count_words(text: str) -> int:
    return len(text.split()) if text else 0
